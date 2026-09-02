"""
Clean the built ReuseHub report: remove every subsection that is not applicable
to this software-only project, producing a final document ready to print.

Removals (identified by their current indices in the filled document):
  - Section 1.7 "Project-Type-Specific Implementation" and its three inapplicable
    subsections (Hardware/IoT, AI/ML, Network).
  - Sections 2.7.1, 2.7.2 and 2.7.3 (AI, IoT, Networking results); the
    applicable 2.7.4 Web/Mobile subsection is retained.
  - Appendix H "Calibration or Performance Records" (not applicable).

Paragraph object references are captured before removal so index shifting does
not matter; every block is deleted as (heading, guidance, placeholder).
"""

from pathlib import Path

import docx

BASE = Path("/home/kim-lee/Desktop/e-mall")
SOURCE = BASE / "ReuseHub_Final_Report_2025-2026.docx"
OUTPUT = BASE / "ReuseHub_Final_Report_FINAL_2025-2026.docx"

doc = docx.Document(str(SOURCE))
paras = doc.paragraphs

# Current (post-build) paragraph indices to remove, grouped by block of 3.
# Each block = [Heading, Template Instruction/Guidance, Student Placeholder].
REMOVE_BLOCKS = [
    # Section 1.7 heading + guidance + 3 inapplicable subsections (11 paragraphs)
    (87, 88, 89, 90, 91, 92, 93, 94, 95, 96, 97),
    # 2.7.1 AI results block (Heading3 + guidance + placeholder)
    (133, 134, 135),
    # 2.7.2 IoT results block
    (136, 137, 138),
    # 2.7.3 Networking results block
    (139, 140, 141),
    # Appendix H block
    (225, 226, 227),
]

# Sanity-check that these are indeed the intended blocks before removing.
def expect(idx, text_fragment):
    p = paras[idx]
    assert text_fragment.lower() in p.text.lower(), f"Unexpected [{idx}]: {p.text[:60]!r} (want ~{text_fragment!r})"

# 1.7 section sanity
expect(87, "1.7 Project-Type-Specific Implementation")
expect(89, "1.7.1 Hardware")
expect(92, "1.7.2 AI")
expect(95, "1.7.3 Network")
expect(97, "Not applicable")
# 2.7 sanity
expect(131, "2.7 Project-Specific Technical Results")
expect(133, "2.7.1 AI")
expect(136, "2.7.2 IoT")
expect(139, "2.7.3 Networking")
expect(142, "2.7.4 Web/Mobile")
expect(144, "Web/mobile results")
# Appendix H sanity
expect(225, "Appendix H: Calibration")
expect(227, "Not applicable")

# Capture paragraph element objects to remove (by object, order-independent).
to_remove = []
for block in REMOVE_BLOCKS:
    for idx in block:
        to_remove.append(paras[idx]._element)

for el in to_remove:
    el.getparent().remove(el)

# Remove every "GUIDANCE: ..." template-instruction paragraph. These are
# developer notes baked into the official template and must not appear in the
# printed report.
guidance_removed = 0
for p in doc.paragraphs:
    if p.style.name == "Template Instruction" and p.text.strip().startswith("GUIDANCE"):
        p._element.getparent().remove(p._element)
        guidance_removed += 1
print(f"Removed guidance paragraphs: {guidance_removed}")

# Remove the front-matter instructional headings and their tables:
#   "Document Purpose and Use", "Mandatory Cross-Referencing Rule",
#   "Content Prohibited from Repetition"
# These are internal template instructions, not part of the printed report.
FRONT_HEADINGS = {
    "Document Purpose and Use",
    "Mandatory Cross-Referencing Rule",
    "Content Prohibited from Repetition",
}
for p in doc.paragraphs:
    if p.style.name == "Heading 1" and p.text.strip() in FRONT_HEADINGS:
        p._element.getparent().remove(p._element)
    elif p.style.name == "Heading 2" and p.text.strip() in FRONT_HEADINGS:
        p._element.getparent().remove(p._element)

# Remove the two front-matter instruction tables (by their header content).
def _header0(t):
    return " ".join(c.text.strip().lower() for c in t.rows[0].cells)

removed_tables = 0
for t in list(doc.tables):
    h = _header0(t)
    if ("evidence provided" in h) or ("do not repeat" in h):
        t._element.getparent().remove(t._element)
        removed_tables += 1
print(f"Removed front-matter instruction tables: {removed_tables}")

# Clean up now-empty paragraphs left in the front matter (blank stray lines).
for p in list(doc.paragraphs):
    if p.text.strip() == "" and p.style.name == "Normal":
        # Only remove blank paragraphs between the removed headings region to
        # avoid over-deleting. We keep it conservative: skip removal here.
        pass

# Re-run the 2.7 placeholders so numbering reads 2.7.1 = Web/Mobile? The template
# labeled it 2.7.4. Since we removed 2.7.1-2.7.3, renumber 2.7.4 -> 2.7.1 for
# clean sequential numbering.
for p in paras:
    if p.style.name == "Heading 3" and p.text.startswith("2.7.4"):
        # replace text runs
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        p.add_run("2.7.1 Web/Mobile/Desktop Results")

doc.save(str(OUTPUT))
print(f"Saved: {OUTPUT}")
